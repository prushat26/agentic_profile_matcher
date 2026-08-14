"""
fs_tools.py - Core File System Tools for LLM Function Calling Assignment
"""

from datetime import datetime
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

# Third-party imports for document parsing
import docx
from pypdf import PdfReader


# ==========================================
# HELPER FUNCTIONS
# ==========================================

def _get_file_metadata(filepath: Path) -> Dict[str, Any]:
    """Extract OS-level file metadata."""
    stat = filepath.stat()
    return {
        "filename": filepath.name,
        "filepath": str(filepath.resolve()),
        "file_extension": filepath.suffix.lower(),
        "size_bytes": stat.st_size,
        "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


def _read_pdf(filepath: Path) -> str:
    """Extract text from a PDF document."""
    reader = PdfReader(str(filepath))
    text_content = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_content.append(page_text)
    return "\n".join(text_content)


def _read_docx(filepath: Path) -> str:
    """Extract text from a DOCX document (paragraphs & tables)."""
    doc = docx.Document(str(filepath))
    text_content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))

    return "\n".join(text_content)


def _read_txt(filepath: Path) -> str:
    """Extract text from a TXT document with encoding fallback."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(filepath, "r", encoding="latin-1") as f:
            return f.read()


# ==========================================
# ASSIGNMENT DELIVERABLE TOOLS
# ==========================================

def read_file(filepath: str) -> Dict[str, Any]:
    """Read resume files (PDF, TXT, DOCX) and return content with metadata."""
    path_obj = Path(filepath)

    if not path_obj.exists():
        return {
            "status": "error",
            "error_type": "FileNotFoundError",
            "message": f"File not found: '{filepath}'",
            "content": None,
            "metadata": None,
        }

    if not path_obj.is_file():
        return {
            "status": "error",
            "error_type": "NotAFileError",
            "message": f"Path exists but is not a file: '{filepath}'",
            "content": None,
            "metadata": None,
        }

    try:
        metadata = _get_file_metadata(path_obj)
        ext = path_obj.suffix.lower()

        if ext == ".pdf":
            content = _read_pdf(path_obj)
        elif ext == ".docx":
            content = _read_docx(path_obj)
        elif ext in [".txt", ".md", ".log", ".csv"]:
            content = _read_txt(path_obj)
        else:
            return {
                "status": "error",
                "error_type": "UnsupportedFormatError",
                "message": f"Unsupported file extension '{ext}'. Supported: .pdf, .docx, .txt",
                "content": None,
                "metadata": metadata,
            }

        return {
            "status": "success",
            "content": content,
            "metadata": metadata,
        }

    except Exception as exc:
        return {
            "status": "error",
            "error_type": exc.__class__.__name__,
            "message": f"Failed to read file: {str(exc)}",
            "content": None,
            "metadata": _get_file_metadata(path_obj) if path_obj.exists() else None,
        }


def list_files(directory: str, extension: Optional[str] = None) -> List[Dict[str, Any]]:
    """List all files in a directory, optionally filtered by extension."""
    dir_path = Path(directory)

    if not dir_path.exists() or not dir_path.is_dir():
        return []

    results = []
    target_ext = extension.lower() if extension else None
    if target_ext and not target_ext.startswith("."):
        target_ext = f".{target_ext}"

    for entry in dir_path.iterdir():
        if entry.is_file():
            if target_ext is None or entry.suffix.lower() == target_ext:
                results.append(_get_file_metadata(entry))

    return results


def write_file(filepath: str, content: str) -> Dict[str, Any]:
    """Write content to file, automatically creating parent directories if needed."""
    path_obj = Path(filepath)

    try:
        # Create directories if needed
        path_obj.parent.mkdir(parents=True, exist_ok=True)

        with open(path_obj, "w", encoding="utf-8") as f:
            f.write(content)

        return {
            "status": "success",
            "filepath": str(path_obj.resolve()),
            "bytes_written": len(content.encode("utf-8")),
            "metadata": _get_file_metadata(path_obj),
        }
    except Exception as exc:
        return {
            "status": "error",
            "error_type": exc.__class__.__name__,
            "message": f"Failed to write file: {str(exc)}",
        }


def search_in_file(filepath: str, keyword: str) -> Dict[str, Any]:
    """Search for keywords in file content and return matches with line context."""
    read_result = read_file(filepath)

    if read_result["status"] != "success":
        return read_result

    content = read_result["content"]
    lines = content.splitlines()
    matches = []
    lower_keyword = keyword.lower()

    for line_num, line in enumerate(lines, start=1):
        if lower_keyword in line.lower():
            matches.append({
                "line_number": line_num,
                "context": line.strip(),
            })

    return {
        "status": "success",
        "filepath": filepath,
        "keyword": keyword,
        "match_count": len(matches),
        "matches": matches,
    }